"""
document_loader.py - Robust Document Ingestion Engine for RAG Pipelines
Supports PDF (PyMuPDF), DOCX (python-docx), and TXT formats.
"""

from dataclasses import dataclass, field
from datetime import datetime, timezone
import hashlib
import re
from pathlib import Path
from typing import Dict, Any, List, Union, Optional, BinaryIO

import fitz  # PyMuPDF
from docx import Document as DocxDocument


@dataclass
class Document:
    """Standardized Document container for RAG processing."""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    doc_id: str = field(init=False)

    def __post_init__(self):
        # Generate deterministic document ID based on clean content (Idempotency)
        self.doc_id = hashlib.sha256(self.content.encode('utf-8')).hexdigest()
        self.metadata["doc_id"] = self.doc_id
        self.metadata["char_count"] = len(self.content)
        self.metadata["word_count"] = len(self.content.split())


class TextCleaner:
    """Text normalization and cleaning utility."""

    # Map Unicode ligatures to standard ascii representation
    LIGATURE_MAP = {
        'ﬁ': 'fi', 'ﬂ': 'fl', 'æ': 'ae', 'œ': 'oe',
        'ﬀ': 'ff', 'ﬃ': 'ffi', 'ﬄ': 'ffl'
    }

    @classmethod
    def clean(cls, text: str) -> str:
        if not text:
            return ""

        # 1. Replace Unicode ligatures
        for lig, repl in cls.LIGATURE_MAP.items():
            text = text.replace(lig, repl)

        # 2. Fix hyphenation at line breaks (e.g., "infor-\nmation" -> "information")
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)

        # 3. Strip non-printable / control characters (keep standard whitespace)
        text = re.sub(r'[^\x20-\x7E\n\t]', '', text)

        # 4. Normalize spaces and repeated empty lines
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n\s*\n+', '\n\n', text)

        return text.strip()


class DocumentLoader:
    """Unified document loader supporting PDF, DOCX, and TXT files."""

    def __init__(self):
        self.cleaner = TextCleaner()

    def load(self, source: Union[str, Path, bytes], file_type: Optional[str] = None, filename: str = "unknown") -> List[Document]:
        """
        Main entry point for loading documents.
        
        Args:
            source: File path (str/Path) or raw bytes.
            file_type: File extension ('pdf', 'docx', 'txt'). Required if passing bytes.
            filename: Display name for metadata tracking.
        """
        if isinstance(source, (str, Path)):
            path = Path(source)
            filename = path.name
            file_type = file_type or path.suffix.lstrip('.').lower()
            with open(path, 'rb') as f:
                raw_bytes = f.read()
        elif isinstance(source, bytes):
            raw_bytes = source
            if not file_type:
                raise ValueError("file_type parameter is required when passing raw bytes.")
        else:
            raise TypeError("Unsupported source type. Must be file path or bytes.")

        file_type = file_type.lower()
        if file_type == 'pdf':
            return self._parse_pdf(raw_bytes, filename)
        elif file_type == 'docx':
            return self._parse_docx(raw_bytes, filename)
        elif file_type == 'txt':
            return self._parse_txt(raw_bytes, filename)
        else:
            raise ValueError(f"Unsupported file format: '.{file_type}'. Supported: pdf, docx, txt.")

    def _parse_pdf(self, raw_bytes: bytes, filename: str) -> List[Document]:
        """Extracts text page-by-page using PyMuPDF (fitz)."""
        documents = []
        # Open PDF directly from bytes (ideal for FastAPI/web endpoints)
        doc = fitz.open(stream=raw_bytes, filetype="pdf")

        for page_num in range(len(doc)):
            page = doc[page_num]
            raw_text = page.get_text("text")
            cleaned_text = self.cleaner.clean(raw_text)

            if not cleaned_text:
                continue  # Skip blank or non-text (image-only) pages

            metadata = {
                "source": filename,
                "page_number": page_num + 1,
                "total_pages": len(doc),
                "doc_type": "pdf",
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            documents.append(Document(content=cleaned_text, metadata=metadata))

        doc.close()
        return documents

    def _parse_docx(self, raw_bytes: bytes, filename: str) -> List[Document]:
        """Extracts text from DOCX files, identifying headings for section context."""
        import io
        doc = DocxDocument(io.BytesIO(raw_bytes))
        documents = []
        current_section = "Header/Introduction"
        buffer = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            # Detect document section headers based on paragraph style
            if p.style.name.startswith("Heading"):
                # Save prior section content
                if buffer:
                    cleaned_content = self.cleaner.clean("\n".join(buffer))
                    if cleaned_content:
                        metadata = {
                            "source": filename,
                            "section_title": current_section,
                            "doc_type": "docx",
                            "created_at": datetime.now(timezone.utc).isoformat()
                        }
                        documents.append(Document(content=cleaned_content, metadata=metadata))
                    buffer.clear()
                current_section = text
            else:
                buffer.append(text)

        # Flush remaining paragraph buffer
        if buffer:
            cleaned_content = self.cleaner.clean("\n".join(buffer))
            if cleaned_content:
                metadata = {
                    "source": filename,
                    "section_title": current_section,
                    "doc_type": "docx",
                    "created_at": datetime.now(timezone.utc).isoformat()
                }
                documents.append(Document(content=cleaned_content, metadata=metadata))

        return documents

    def _parse_txt(self, raw_bytes: bytes, filename: str) -> List[Document]:
        """Parses plain text files."""
        text = raw_bytes.decode('utf-8', errors='ignore')
        cleaned_text = self.cleaner.clean(text)

        if not cleaned_text:
            return []

        metadata = {
            "source": filename,
            "doc_type": "txt",
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        return [Document(content=cleaned_text, metadata=metadata)]


# --- Verification Example ---
if __name__ == "__main__":
    loader = DocumentLoader()

    # 1. Verification with dummy sample TXT
    sample_txt_bytes = b"Hello RAG World!\n\nThis is a test  document with   ligature \x00 characters like \xef\xac\x81."
    docs = loader.load(sample_txt_bytes, file_type="txt", filename="test_sample.txt")

    print(f"Loaded {len(docs)} document(s):")
    for d in docs:
        print("--- Document Output ---")
        print(f"Doc ID (SHA256): {d.doc_id}")
        print(f"Content: {d.content}")
        print(f"Metadata: {d.metadata}")